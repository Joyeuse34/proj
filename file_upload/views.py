import pandas as pd 
from docx import Document
import os 
import zipfile 
from django.shortcuts import render 
from django.http import HttpResponse 
from .models import FileUpload 
from .forms import FileUploadForm 
from django.conf import settings
from datetime import datetime

def replace_placeholder(doc, placeholder, replacement): 
    for paragraph in doc.paragraphs: 
        if placeholder in paragraph.text: 
            paragraph.text = paragraph.text.replace(placeholder, str(replacement))
            for table in doc.tables:
                for row in table.rows: 
                    for cell in row.cells: 
                        if placeholder in cell.text: cell.text = cell.text.replace(placeholder, str(replacement))

def upload_and_process(request): 
    if request.method == 'POST': 
        form = FileUploadForm(request.POST, request.FILES) 
        if form.is_valid():  
            file_upload = form.save() 
            excel_file = file_upload.excel_file.path 
            yollanma_template = file_upload.yollanma_template.path
            shartnoma_template = file_upload.shartnoma_template.path
            kundalik_template = file_upload.kundalik_template.path
            current_date = datetime.now().strftime("%d.%m.%Y")
            df = pd.read_excel(excel_file)
            output_dir = os.path.join(settings.MEDIA_ROOT, 'output')
            os.makedirs(output_dir, exist_ok=True)
            output_files = []
            for index, row in df.iterrows():
                student_id = [row["№"]]
                doc_yollanma = Document(yollanma_template)
                doc_shartnoma = Document(shartnoma_template)
                doc_kundalik = Document(kundalik_template)
                replacements = {
                    '<student>': row['F.I.SH'],  
                    '<studentName>': row['F.I.SH'],
                    '<companyName>': row['Amaliyot o\'tish joyi'],  
                    '<location>':row['Manzili'],
                    '<facultyHead>': row['Kafedra bo\'yicha rahbar'],  
                    '<internshipHead>': row['Korxonaga ma\'sul shaxs'],  
                    '<date>':current_date,
                    '<id>':str(student_id)
                }
                for placeholder, replacement in replacements.items():
                    replace_placeholder(doc_kundalik, placeholder, replacement)
                    replace_placeholder(doc_shartnoma, placeholder, replacement)
                    replace_placeholder(doc_yollanma, placeholder, replacement)
                output_kundalik = os.path.join(output_dir, f'1.1_Kundalik_{int(row["№"])}.docx')
                output_yollanma = os.path.join(output_dir, f'1.2_Yollanma_{int(row["№"])}.docx')
                output_shartnoma = os.path.join(output_dir, f'1.3_Shartnoma_{int(row["№"])}.docx')
                doc_kundalik.save(output_kundalik)
                doc_shartnoma.save(output_shartnoma)
                doc_yollanma.save(output_yollanma)
                output_files.extend([output_kundalik,output_yollanma,output_shartnoma])
            
        zip_path = os.path.join(settings.MEDIA_ROOT, 'output.zip')
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for output_file in output_files:
                zipf.write(output_file, os.path.basename(output_file))
        
        with open(zip_path, 'rb') as zip_file:
            response = HttpResponse(zip_file.read(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename=output.zip'
            return response
    else:
        form = FileUploadForm()
    return render(request, 'file_upload/upload.html', {'form': form})